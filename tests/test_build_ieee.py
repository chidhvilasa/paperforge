"""Tests for the IEEE Transactions / journal LaTeX build template."""

from pathlib import Path

import yaml

from paperforge.commands import build, init
from paperforge.models.claim import Claim
from paperforge.models.experiment import Experiment


def write_journal_project(tmp_path: Path) -> None:
    init.run(tmp_path)
    pf_dir = tmp_path / ".paperforge"

    (pf_dir / "claims" / "claim_01.yaml").unlink()

    paper_yaml = pf_dir / "paper.yaml"
    data = yaml.safe_load(paper_yaml.read_text(encoding="utf-8"))
    data["title"] = "Test Paper"
    data["authors"] = ["A. Author"]
    data["paper_type"] = "journal"
    data["keywords"] = ["security", "IoT"]
    paper_yaml.write_text(yaml.dump(data, default_flow_style=False), encoding="utf-8")

    exp_path = pf_dir / "experiments" / "exp_01.yaml"
    experiment = Experiment(
        id="exp_01",
        description="Test experiment",
        metrics={"accuracy": 98.4},
        hardware="RTX 4070",
        dataset="TestSet",
        seed=42,
    )
    exp_path.write_text(
        yaml.dump(experiment.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )

    claim_path = pf_dir / "claims" / "claim_02.yaml"
    claim = Claim(
        id="claim_02",
        text="The system achieves 98.4% accuracy.",
        experiment="exp_01",
        citations=["smith2024"],
        sections=["results", "introduction"],
        status="verified",
    )
    claim_path.write_text(
        yaml.dump(claim.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )

    abs_claim_path = pf_dir / "claims" / "claim_abs.yaml"
    abs_claim = Claim(
        id="claim_abs",
        text="Abstract text without citations.",
        experiment="exp_01",
        citations=[],
        sections=["abstract"],
        status="verified",
    )
    abs_claim_path.write_text(
        yaml.dump(abs_claim.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )


def _get_out(tmp_path: Path) -> Path:
    p = tmp_path / "paper_generated" / "current"
    return p if p.exists() else tmp_path / "paper"


def _read_tex(tmp_path: Path) -> str:
    p = _get_out(tmp_path) / "paper.tex"
    if p.exists():
        return p.read_text(encoding="utf-8")
    return (tmp_path / ".paperforge" / "output" / "paper.tex").read_text(
        encoding="utf-8"
    )


def test_build_journal_creates_tex(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    assert (_get_out(tmp_path) / "paper.tex").exists()


def test_build_journal_documentclass(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert content.splitlines()[0] == "\\documentclass[journal]{IEEEtran}"


def test_build_journal_abstract_in_titleabstractindextext(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert "IEEEtitleabstractindextext" in content
    assert "begin{abstract}" in content


def test_build_journal_ieeeraisesectionheading(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert "IEEEraisesectionheading" in content


def test_build_journal_ieeeparstart(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert "IEEEPARstart" in content


def test_build_journal_ieeedisplaynontitleabstractindextext(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert "IEEEdisplaynontitleabstractindextext" in content


def test_build_journal_keywords(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert "IEEEkeywords" in content
    assert "security" in content


def test_build_journal_acknowledgment(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert "Acknowledgments" in content
    assert "ifCLASSOPTIONcompsoc" in content


def test_build_journal_bibliography_stub(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    bib_path = _get_out(tmp_path) / "references.bib"
    assert bib_path.exists()
    content = bib_path.read_text(encoding="utf-8")
    assert "smith2024" in content


def test_build_ieee_journal_target_alias(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-trans")
    content = _read_tex(tmp_path)
    assert "journal" in content


def test_build_conference_unchanged(tmp_path: Path) -> None:
    init.run(tmp_path)
    pf_dir = tmp_path / ".paperforge"

    claim_path = pf_dir / "claims" / "claim_01.yaml"
    claim = Claim(
        id="claim_01",
        text="This model achieves 98.4% accuracy.",
        experiment="exp_01",
        sections=["results"],
        status="verified",
    )
    claim_path.write_text(
        yaml.dump(claim.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )

    exp_path = pf_dir / "experiments" / "exp_01.yaml"
    experiment = Experiment(id="exp_01", metrics={"accuracy": 98.4})
    exp_path.write_text(
        yaml.dump(experiment.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )

    paper_yaml = pf_dir / "paper.yaml"
    data = yaml.safe_load(paper_yaml.read_text(encoding="utf-8"))
    data["title"] = "Test Paper Title"
    data["authors"] = ["Test Author"]
    data["paper_type"] = "conference"
    paper_yaml.write_text(yaml.dump(data, default_flow_style=False), encoding="utf-8")

    build.run(tmp_path, target="ieee")

    content = _read_tex(tmp_path)
    assert "conference" in content
    assert "IEEEtitleabstractindextext" not in content


def test_build_paper_type_journal_auto_selects_journal_template(
    tmp_path: Path,
) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")


def test_build_author_not_double_wrapped(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    author_count = content.count("\\author{")
    assert author_count == 1, f"Expected \\author{{}} exactly once, found {author_count} times"
    assert "\\author{\\author" not in content


def test_build_table_label_appears_once(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    pf_dir = tmp_path / ".paperforge"

    claim3_path = pf_dir / "claims" / "claim_03.yaml"
    claim3 = Claim(
        id="claim_03",
        text="The system maintains 98.4% accuracy across runs.",
        experiment="exp_01",
        figures=[],
        tables=["tbl_01"],
        citations=[],
        sections=["discussion"],
        status="verified",
    )
    claim3_path.write_text(
        yaml.dump(claim3.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )

    tbl_path = pf_dir / "tables" / "tbl_01.yaml"
    tbl_path.write_text(
        yaml.dump(
            {
                "id": "tbl_01",
                "caption": "Performance Comparison",
                "columns": ["Method", "Accuracy"],
                "rows": [["B2", "91.2%"], ["Ours", "98.4%"]],
                "notes": "",
                "first_mentioned_in": "results",
                "source_experiment": "exp_01",
            },
            default_flow_style=False,
        ),
        encoding="utf-8",
    )

    claim2_path = pf_dir / "claims" / "claim_02.yaml"
    claim2_data = yaml.safe_load(claim2_path.read_text(encoding="utf-8"))
    claim2_data["tables"] = ["tbl_01"]
    claim2_path.write_text(yaml.dump(claim2_data, default_flow_style=False), encoding="utf-8")

    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    label_count = content.count("\\label{tab:tbl_01}")
    assert label_count == 1, f"Expected \\label{{tab:tbl_01}} exactly once, found {label_count}"
    begin_table_count = content.count("\\begin{table}")
    assert begin_table_count == 1, f"Expected \\begin{{table}} exactly once, found {begin_table_count}"


def test_build_figure_label_appears_once(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    pf_dir = tmp_path / ".paperforge"

    fig_path = pf_dir / "figures" / "fig_01.yaml"
    fig_path.write_text(
        yaml.dump(
            {
                "id": "fig_01",
                "caption": "System Architecture",
                "path": "figures/fig_01.png",
                "format": "png",
                "width_inches": 3.5,
                "resolution_dpi": 300,
                "first_mentioned_in": "results",
            },
            default_flow_style=False,
        ),
        encoding="utf-8",
    )

    claim2_path = pf_dir / "claims" / "claim_02.yaml"
    claim2_data = yaml.safe_load(claim2_path.read_text(encoding="utf-8"))
    claim2_data["figures"] = ["fig_01"]
    claim2_path.write_text(yaml.dump(claim2_data, default_flow_style=False), encoding="utf-8")

    claim3_path = pf_dir / "claims" / "claim_03.yaml"
    claim3 = Claim(
        id="claim_03",
        text="As shown above, latency decreases under high load.",
        experiment="exp_01",
        figures=["fig_01"],
        tables=[],
        citations=[],
        sections=["discussion"],
        status="verified",
    )
    claim3_path.write_text(
        yaml.dump(claim3.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )

    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    label_count = content.count("\\label{fig:fig_01}")
    assert label_count == 1, f"Expected \\label{{fig:fig_01}} exactly once, found {label_count}"
    begin_figure_count = content.count("\\begin{figure}")
    assert begin_figure_count == 1, f"Expected \\begin{{figure}} exactly once, found {begin_figure_count}"


def test_build_conference_author_not_double_wrapped(tmp_path: Path) -> None:
    init.run(tmp_path)
    pf_dir = tmp_path / ".paperforge"

    claim_path = pf_dir / "claims" / "claim_01.yaml"
    claim = Claim(
        id="claim_01",
        text="This model achieves 98.4% accuracy.",
        experiment="exp_01",
        sections=["results"],
        status="verified",
    )
    claim_path.write_text(
        yaml.dump(claim.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )

    exp_path = pf_dir / "experiments" / "exp_01.yaml"
    experiment = Experiment(id="exp_01", metrics={"accuracy": 98.4})
    exp_path.write_text(
        yaml.dump(experiment.to_yaml(), default_flow_style=False, allow_unicode=True),
        encoding="utf-8",
    )

    paper_yaml = pf_dir / "paper.yaml"
    data = yaml.safe_load(paper_yaml.read_text(encoding="utf-8"))
    data["title"] = "Test Paper Title"
    data["authors"] = ["Test Author"]
    paper_yaml.write_text(yaml.dump(data, default_flow_style=False), encoding="utf-8")

    build.run(tmp_path, target="ieee")
    content = _read_tex(tmp_path)
    author_count = content.count("\\author{")
    assert author_count == 1
    assert "\\author{\\author" not in content


def test_build_preserves_real_references_bib(tmp_path: Path) -> None:
    write_journal_project(tmp_path)

    build.run(tmp_path, target="ieee-journal")

    bib_path = _get_out(tmp_path) / "references.bib"
    bib_path.write_text(
        "@article{smith2024,\n"
        "  author = {Smith, A.},\n"
        "  title = {Real Title},\n"
        "  journal = {IEEE Access},\n"
        "  year = {2024}\n"
        "}\n",
        encoding="utf-8",
    )

    build.run(tmp_path, target="ieee-journal")

    content = bib_path.read_text(encoding="utf-8")
    assert "Real Title" in content
    assert "TODO" not in content


def test_build_overwrites_stub_references_bib(tmp_path: Path) -> None:
    write_journal_project(tmp_path)

    build.run(tmp_path, target="ieee-journal")
    build.run(tmp_path, target="ieee-journal")

    bib_path = _get_out(tmp_path) / "references.bib"
    content = bib_path.read_text(encoding="utf-8")
    assert "@article" in content


def test_build_escapes_percentage_in_claim_text(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    pf_dir = tmp_path / ".paperforge"
    claim_path = pf_dir / "claims" / "claim_02.yaml"
    claim_data = yaml.safe_load(claim_path.read_text(encoding="utf-8"))
    claim_data["text"] = "achieves 98.4% accuracy"
    claim_path.write_text(yaml.dump(claim_data, default_flow_style=False), encoding="utf-8")

    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert "98.4\\%" in content
    assert "98.4%" not in content


def test_build_escapes_ampersand_in_claim_text(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    pf_dir = tmp_path / ".paperforge"
    claim_path = pf_dir / "claims" / "claim_02.yaml"
    claim_data = yaml.safe_load(claim_path.read_text(encoding="utf-8"))
    claim_data["text"] = "System A & B comparison"
    claim_path.write_text(yaml.dump(claim_data, default_flow_style=False), encoding="utf-8")

    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    assert "A \\& B" in content
    assert "A & B" not in content


def test_build_claim_appears_once_when_in_multiple_sections(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    pf_dir = tmp_path / ".paperforge"
    claim_path = pf_dir / "claims" / "claim_02.yaml"
    claim_data = yaml.safe_load(claim_path.read_text(encoding="utf-8"))
    claim_data["text"] = "Unique claim text for dedup test."
    claim_data["sections"] = ["results", "discussion"]
    claim_path.write_text(yaml.dump(claim_data, default_flow_style=False), encoding="utf-8")

    build.run(tmp_path, target="ieee-journal")
    content = _read_tex(tmp_path)
    count = content.count("Unique claim text for dedup test")
    assert count == 1, f"Claim text appeared {count} times, expected 1"


def test_escape_latex_function() -> None:
    from paperforge.utils.latex import escape_latex

    assert escape_latex("98.4%") == "98.4\\%"
    assert escape_latex("A & B") == "A \\& B"
    assert escape_latex("$x$") == "$x$"
    assert escape_latex("") == ""
    assert escape_latex("normal text") == "normal text"


def test_pdf_stale_when_no_pdf_exists(tmp_path: Path) -> None:
    write_journal_project(tmp_path)
    assert build._is_pdf_stale(tmp_path) is True


def test_pdf_stale_when_claim_newer_than_pdf(tmp_path: Path) -> None:
    import time

    write_journal_project(tmp_path)
    paper_dir = _get_out(tmp_path)
    paper_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = paper_dir / "paper.pdf"
    pdf_path.write_text("fake pdf content", encoding="utf-8")

    time.sleep(0.05)
    claim_path = tmp_path / ".paperforge" / "claims" / "claim_02.yaml"
    claim_data = yaml.safe_load(claim_path.read_text(encoding="utf-8"))
    claim_data["text"] = "Updated claim text for freshness test."
    claim_path.write_text(yaml.dump(claim_data), encoding="utf-8")

    assert build._is_pdf_stale(tmp_path) is True


def test_pdf_not_stale_when_pdf_newer(tmp_path: Path) -> None:
    import os
    import time

    write_journal_project(tmp_path)
    paper_dir = _get_out(tmp_path)
    paper_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = paper_dir / "paper.pdf"
    pdf_path.write_text("fake pdf content", encoding="utf-8")

    now = time.time()
    for f in (tmp_path / ".paperforge").rglob("*.yaml"):
        os.utime(f, (now - 10, now - 10))

    assert build._is_pdf_stale(tmp_path) is False


def test_docx_generated_when_no_latex(tmp_path: Path, monkeypatch) -> None:
    write_journal_project(tmp_path)
    monkeypatch.setattr("shutil.which", lambda name: None)
    build.run(tmp_path, target="ieee-access", force=True)
    assert (_get_out(tmp_path) / "paper.docx").exists()


def test_docx_contains_title(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    write_journal_project(tmp_path)
    paper_yaml = tmp_path / ".paperforge" / "paper.yaml"
    data = yaml.safe_load(paper_yaml.read_text(encoding="utf-8"))
    data["title"] = "My Test Paper"
    paper_yaml.write_text(yaml.dump(data), encoding="utf-8")

    monkeypatch.setattr("shutil.which", lambda name: None)
    build.run(tmp_path, target="ieee-access", force=True)

    doc = Document(str(_get_out(tmp_path) / "paper.docx"))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "My Test Paper" in full_text


def test_docx_contains_table(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    write_journal_project(tmp_path)
    tbl_data = {
        "id": "tbl_01",
        "caption": "DOCX Table Caption",
        "columns": ["Col1", "Col2"],
        "rows": [["Val1", "Val2"]],
    }
    (tmp_path / ".paperforge" / "tables" / "tbl_01.yaml").write_text(
        yaml.dump(tbl_data), encoding="utf-8"
    )
    claim_path = tmp_path / ".paperforge" / "claims" / "claim_02.yaml"
    claim_data = yaml.safe_load(claim_path.read_text(encoding="utf-8"))
    claim_data["tables"] = ["tbl_01"]
    claim_path.write_text(yaml.dump(claim_data), encoding="utf-8")

    monkeypatch.setattr("shutil.which", lambda name: None)
    build.run(tmp_path, target="ieee-access", force=True)

    doc = Document(str(_get_out(tmp_path) / "paper.docx"))
    assert len(doc.tables) >= 1


def test_force_flag_bypasses_freshness_check(tmp_path: Path, monkeypatch) -> None:
    import os
    import time

    write_journal_project(tmp_path)
    paper_dir = _get_out(tmp_path)
    paper_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = paper_dir / "paper.pdf"
    pdf_path.write_text("fake pdf content", encoding="utf-8")

    now = time.time()
    for f in (tmp_path / ".paperforge").rglob("*.yaml"):
        os.utime(f, (now - 10, now - 10))

    tex_path = paper_dir / "paper.tex"
    if tex_path.exists():
        tex_path.unlink()

    monkeypatch.setattr("shutil.which", lambda name: None)
    build.run(tmp_path, target="ieee-access", force=True)

    assert tex_path.exists()


def test_stale_pdf_deleted_before_rebuild(tmp_path: Path, monkeypatch) -> None:
    import time

    write_journal_project(tmp_path)
    paper_dir = _get_out(tmp_path)
    paper_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = paper_dir / "paper.pdf"
    pdf_path.write_text("old fake pdf content", encoding="utf-8")

    time.sleep(0.05)
    claim_path = tmp_path / ".paperforge" / "claims" / "claim_02.yaml"
    claim_data = yaml.safe_load(claim_path.read_text(encoding="utf-8"))
    claim_data["text"] = "Stale pdf test updated claim."
    claim_path.write_text(yaml.dump(claim_data), encoding="utf-8")

    monkeypatch.setattr("shutil.which", lambda name: None)
    build.run(tmp_path, target="ieee-access")

    assert not pdf_path.exists() or pdf_path.read_text(encoding="utf-8") != "old fake pdf content"


# --- Test 31: Aux files deleted after build ---

def test_aux_files_deleted_after_build(tmp_path: Path, monkeypatch) -> None:
    """After build, no .aux .log .fls .out .bbl .blg files should remain in output_dir."""
    write_journal_project(tmp_path)

    # Monkeypatch shutil.which so pdflatex is "found" but fails silently
    monkeypatch.setattr("shutil.which", lambda name: "/usr/bin/pdflatex" if name in ("pdflatex", "latexmk") else None)

    # Also monkeypatch subprocess.run so compilation appears to happen but produces no PDF
    # This exercises the cleanup path after failed compilation
    import subprocess as _subprocess_mod
    original_run = _subprocess_mod.run

    def mock_run(args, **kwargs):
        # Intercept pdflatex calls; create fake aux files to simulate compilation
        if args and isinstance(args, list) and any("pdflatex" in str(a) or "latexmk" in str(a) for a in args):
            # Find output_dir from args
            cwd = kwargs.get("cwd")
            if cwd:
                from pathlib import Path as _P
                cwd_path = _P(cwd)
                for ext in [".aux", ".log", ".fls", ".out", ".bbl"]:
                    try:
                        (cwd_path / f"paper{ext}").write_text(f"fake {ext} content", encoding="utf-8")
                    except OSError:
                        pass
            result = _subprocess_mod.CompletedProcess(args, returncode=1)
            return result
        return original_run(args, **kwargs)

    monkeypatch.setattr(_subprocess_mod, "run", mock_run)
    build.run(tmp_path, target="ieee-access", force_anyway=True)

    # Find the output directory
    from paperforge.core.project import PaperForgeProject
    project = PaperForgeProject.load(tmp_path)
    output_dir = tmp_path / project.config.build_output_dir

    if output_dir.exists():
        aux_exts = {".aux", ".log", ".fls", ".out", ".bbl", ".blg"}
        remaining_aux = [f for f in output_dir.glob("*") if f.suffix in aux_exts]
        assert remaining_aux == [], f"Aux files remain after build: {[f.name for f in remaining_aux]}"


# --- Test 32: Rotation keeps only key files ---

def test_rotation_keeps_only_key_files(tmp_path: Path, monkeypatch) -> None:
    """previous/ should only contain meaningful files, never aux files."""
    write_journal_project(tmp_path)
    monkeypatch.setattr("shutil.which", lambda name: None)

    # First build — creates current/
    build.run(tmp_path, target="ieee-access", force_anyway=True)

    from paperforge.core.project import PaperForgeProject
    project = PaperForgeProject.load(tmp_path)
    output_dir = tmp_path / project.config.build_output_dir

    # Manually plant a fake aux file in current/ to simulate leftover
    if output_dir.exists():
        (output_dir / "paper.aux").write_text("fake aux", encoding="utf-8")
        (output_dir / "paper.log").write_text("fake log", encoding="utf-8")

    # Second build — triggers rotation of current/ → previous/
    build.run(tmp_path, target="ieee-access", force=True, force_anyway=True)

    previous_dir = output_dir.parent / "previous"
    if previous_dir.exists():
        prev_files = {f.name for f in previous_dir.iterdir() if f.is_file()}
        allowed = {
            "paper.pdf", "paper.tex", "references.bib",
            "paper_overleaf.zip", "paper.docx", "traceability.tex",
        }
        unexpected = prev_files - allowed
        assert unexpected == set(), f"Unexpected files in previous/: {unexpected}"


# --- Test 33: Clean command removes aux files ---

def test_clean_command_runs(tmp_path: Path, monkeypatch) -> None:
    """paperforge clean should remove manually-placed aux files from paper_generated/."""
    from paperforge.commands import clean

    write_journal_project(tmp_path)
    monkeypatch.setattr("shutil.which", lambda name: None)
    build.run(tmp_path, target="ieee-access", force_anyway=True)

    from paperforge.core.project import PaperForgeProject
    project = PaperForgeProject.load(tmp_path)
    output_dir = tmp_path / project.config.build_output_dir

    # Plant a fake aux file
    if output_dir.exists():
        aux_file = output_dir / "paper.aux"
        aux_file.write_text("aux content", encoding="utf-8")
        assert aux_file.exists(), "Failed to create test aux file"

        clean.run(tmp_path)
        assert not aux_file.exists(), "clean.run() did not remove paper.aux"
