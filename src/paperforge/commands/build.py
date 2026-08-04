"""paperforge build command."""

from __future__ import annotations

import json
import platform
import re
import shutil
import subprocess
import sys
from datetime import UTC
from pathlib import Path
from typing import Any

from rich.console import Console, Group
from rich.panel import Panel
from rich.text import Text

from paperforge.commands.doctor import collect_issues
from paperforge.core.project import (
    Affiliation,
    Biography,
    PaperForgeProject,
    ProjectConfig,
)
from paperforge.models.claim import Claim
from paperforge.models.figure import Figure
from paperforge.models.table import Table
from paperforge.utils.latex import (
    escape_latex,
    escape_latex_safe,
    markdown_to_latex_inline,
)
from paperforge.venues.base import VenuePlugin
from paperforge.venues.registry import get_plugin

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

console = Console()

SECTION_TITLES = {
    "introduction": "Introduction",
    "related_work": "Related Work",
    "methodology": "Methodology",
    "experiments": "Experimental Setup",
    "results": "Results",
    "discussion": "Discussion",
    "conclusion": "Conclusion",
}


def _claim_paragraph(claim: Claim, project: PaperForgeProject) -> str:
    raw = claim.is_math or claim.raw_latex
    text_to_emit = escape_latex_safe(claim.text, raw=raw)
    if not raw:
        text_to_emit = markdown_to_latex_inline(text_to_emit)

    if claim.citations:
        cite_block = "\\cite{" + ",".join(claim.citations) + "}"
        text = text_to_emit.rstrip()
        if text.endswith("."):
            paragraph = text[:-1] + f" {cite_block}."
        else:
            paragraph = text + f" {cite_block}."
    else:
        paragraph = text_to_emit

    first_figure_yaml = None
    for figure_id in claim.figures:
        fig_obj = next((f for f in project.figures if f.id == figure_id), None)
        if fig_obj and not first_figure_yaml:
            first_figure_yaml = fig_obj

    refs = []
    if first_figure_yaml:
        refs.append(f"Fig.~\\ref{{fig:{first_figure_yaml.id}}}")
    for table in claim.tables:
        refs.append(f"Table~\\ref{{tab:{table}}}")

    if refs:
        paragraph += (
            f" (see {', and '.join(refs) if len(refs) > 2 else ' and '.join(refs)})"
        )

    if claim.claim_type == "proof":
        paragraph = f"\\begin{{proof}}\n{paragraph}\n\\end{{proof}}"
    elif claim.claim_type in ("theorem", "lemma", "definition", "corollary", "remark"):
        label = f"{claim.claim_type}:{claim.id}"
        paragraph = f"\\begin{{{claim.claim_type}}}\n\\label{{{label}}}\n{paragraph}\n\\end{{{claim.claim_type}}}"

    return paragraph


def _generate_abstract(claims: list[Claim]) -> str:
    abstract_claims = [c for c in claims if "abstract" in c.sections]
    if not abstract_claims:
        return "% TODO: Add claims to the abstract section."
    parts = []
    for c in sorted(abstract_claims, key=lambda claim: claim.id):
        raw = c.is_math or c.raw_latex
        txt = escape_latex_safe(c.text, raw=raw)
        if not raw:
            txt = markdown_to_latex_inline(txt)
        parts.append(txt)
    return " ".join(parts)


def _normalize_asset_rel_path(raw_path: str) -> str | None:
    """Normalize a user-configured asset path to a safe, portable,
    forward-slash relative path.

    Returns None if the path is empty, absolute (POSIX or Windows
    drive-letter), or escapes its base directory via '..'.
    """
    normalized = raw_path.replace("\\", "/").strip()
    if not normalized:
        return None
    if normalized.startswith("/"):
        return None
    if len(normalized) > 1 and normalized[1] == ":":
        return None
    parts = [p for p in normalized.split("/") if p not in ("", ".")]
    if not parts or any(p == ".." for p in parts):
        return None
    return "/".join(parts)


def resolve_figure_asset(
    fig_obj: Figure, project_root: Path, output_dir: Path | None = None
) -> str | None:
    """Resolve a figure's configured path against the project root and,
    if an output directory is given, materialize (copy) the asset into it
    so that generated LaTeX can reference it with a package/output-relative
    path -- never a path relative to project_root.

    Returns the LaTeX-relative asset path (forward-slash, portable) if the
    source asset was found and (when output_dir is given) successfully
    copied; returns None if the source could not be resolved.
    """
    if not fig_obj.path:
        return None
    rel_path = _normalize_asset_rel_path(fig_obj.path)
    if rel_path is None:
        return None
    source = project_root / rel_path
    try:
        if not source.is_file():
            return None
    except OSError:
        return None
    if output_dir is None:
        return rel_path
    dest = output_dir / rel_path
    try:
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(source), str(dest))
    except OSError:
        return None
    return rel_path


def _generate_figure_latex(
    fig_obj: Figure,
    project_root: Path | None = None,
    output_dir: Path | None = None,
) -> str:
    env = "figure*" if fig_obj.wide else "figure"
    escaped_caption = escape_latex_safe(fig_obj.caption or "", raw=fig_obj.is_math)

    resolved_path: str | None = None
    if project_root is not None:
        resolved_path = resolve_figure_asset(fig_obj, project_root, output_dir)

    if resolved_path:
        if fig_obj.width_inches:
            width = f"{fig_obj.width_inches}in"
        elif fig_obj.wide:
            width = "\\textwidth"
        else:
            width = "\\columnwidth"
        path = resolved_path
        return (
            f"\\begin{{{env}}}[!t]\n"
            f"\\centering\n"
            f"\\includegraphics[width={width}]{{{path}}}\n"
            f"\\caption{{{escaped_caption}}}\n"
            f"\\label{{fig:{fig_obj.id}}}\n"
            f"\\end{{{env}}}"
        )

    # Missing figure placeholder box
    path_display = escape_latex_safe(fig_obj.path or "not set")
    return (
        f"\\begin{{{env}}}[!t]\n"
        f"\\centering\n"
        f"\\fbox{{\\parbox{{0.9\\columnwidth}}{{\\centering\n"
        f"\\textbf{{Figure placeholder}}\\\\\n"
        f"{escaped_caption}\\\\\n"
        f"Path: {path_display}\n"
        f"}}}}\n"
        f"\\caption{{{escaped_caption}}}\n"
        f"\\label{{fig:{fig_obj.id}}}\n"
        f"\\end{{{env}}}"
    )


def _generate_table_latex(
    table: Table, project: PaperForgeProject | None = None
) -> str:
    env = "table*" if table.wide else "table"
    escaped_caption = escape_latex_safe(table.caption or "", raw=table.is_math)

    columns = list(table.columns)
    rows = list(table.rows)

    if not rows and table.auto_rows_from_experiment and project:
        exp_obj = next(
            (e for e in project.experiments if e.id == table.auto_rows_from_experiment),
            None,
        )
        if exp_obj and exp_obj.metrics:
            if not columns:
                columns = ["Metric", "Value"]
            rows = [[m_name, str(m_val)] for m_name, m_val in exp_obj.metrics.items()]

    if not columns:
        # No columns defined -- emit a comment placeholder
        return (
            f"% Table: {table.id} — {escaped_caption[:60]}\n"
            f"% (no column data — fill in .paperforge/tables/{table.id}.yaml)\n"
            f"% \\label{{tab:{table.id}}}"
        )

    col_spec = " ".join(["c"] * len(columns))
    header_row = (
        " & ".join(escape_latex_safe(c, raw=table.raw_latex_rows) for c in columns)
        + " \\\\"
    )

    data_rows = []
    for row in rows:
        padded = row[: len(columns)]
        while len(padded) < len(columns):
            padded.append("")
        escaped_row = [
            escape_latex_safe(cell, raw=table.raw_latex_rows) for cell in padded
        ]
        data_rows.append(" & ".join(escaped_row) + " \\\\")

    notes_block = ""
    if table.notes:
        notes_block = f"\n\\footnotesize{{\\textit{{Note: {escape_latex_safe(table.notes, raw=table.is_math)}}}}}"

    lines = [
        f"\\begin{{{env}}}[!t]",
        "\\renewcommand{\\arraystretch}{1.3}",
        f"\\caption{{{escaped_caption}}}",
        f"\\label{{tab:{table.id}}}",
        "\\centering",
        f"\\begin{{tabular}}{{{col_spec}}}",
        "\\hline",
        header_row,
        "\\hline",
    ]
    lines.extend(data_rows)
    lines.append("\\hline")
    lines.append("\\end{tabular}")
    if notes_block:
        lines.append(notes_block)
    lines.append(f"\\end{{{env}}}")

    return "\n".join(lines)


def _generate_sections_overview(overview: str) -> str:
    if not overview or not overview.strip():
        return ""
    text = overview.strip()
    standalone_prefixes = (
        "section",
        "the rest",
        "this paper",
        "the remainder",
        "the paper",
    )
    if text.lower().startswith(standalone_prefixes):
        return escape_latex(text)
    return f"The rest of this paper is organized as follows: {escape_latex(text)}"


def _generate_sections(
    sections: list[str], project: PaperForgeProject, output_dir: Path | None = None
) -> str:
    blocks: list[str] = []
    emitted_figures: set[str] = set()
    emitted_tables: set[str] = set()
    emitted_algorithms: set[str] = set()
    emitted_claims: set[str] = set()

    for section in sections:
        if section == "abstract":
            continue
        title = SECTION_TITLES.get(section, section.replace("_", " ").title())
        section_claims = sorted(
            (c for c in project.claims if section in c.sections), key=lambda c: c.id
        )
        block = f"\\section{{{title}}}\n"
        current_subsection = ""

        if section_claims:
            claim_blocks = []
            non_contrib = [c for c in section_claims if not c.is_contribution]
            contrib = [c for c in section_claims if c.is_contribution]
            target_claims = (
                non_contrib
                if (section == "introduction" and contrib)
                else section_claims
            )

            for c in target_claims:
                if c.id in emitted_claims:
                    first_sec = c.sections[0] if c.sections else "above"
                    claim_blocks.append(
                        f"% Claim {c.id} content appears in {first_sec} above."
                    )
                    continue
                emitted_claims.add(c.id)

                prefix = ""
                if c.subsection and c.subsection != current_subsection:
                    prefix = f"\\subsection{{{escape_latex(c.subsection)}}}\n"
                    current_subsection = c.subsection

                if section == "related_work" and c.compared_work:
                    prefix += f"% --- {escape_latex(c.compared_work)} ---\n"

                text_par = prefix + _claim_paragraph(c, project)

                fig_envs = []
                for fig_id in c.figures:
                    fig_obj = next((f for f in project.figures if f.id == fig_id), None)
                    if fig_obj and fig_id not in emitted_figures:
                        emitted_figures.add(fig_id)
                        fig_envs.append(
                            _generate_figure_latex(fig_obj, project.root, output_dir)
                        )
                    elif fig_obj and fig_id in emitted_figures:
                        fig_envs.append(
                            f"% Figure~\\ref{{fig:{fig_id}}} already defined above."
                        )
                    else:
                        fig_envs.append(
                            f"% Reference: {fig_id} (no figure YAML — run paperforge add-figure)"
                        )
                if fig_envs:
                    text_par += "\n\n" + "\n\n".join(fig_envs)

                tbl_envs = []
                for tbl_id in c.tables:
                    tbl_obj = next((t for t in project.tables if t.id == tbl_id), None)
                    if tbl_obj and tbl_id not in emitted_tables:
                        emitted_tables.add(tbl_id)
                        if tbl_obj.caption or tbl_obj.auto_rows_from_experiment:
                            tbl_envs.append(_generate_table_latex(tbl_obj, project))
                        else:
                            caption_text = escape_latex((tbl_obj.caption or "")[:60])
                            tbl_envs.append(
                                f"% Table: {tbl_id} — {caption_text} (no caption set)\n"
                                f"% \\label{{tab:{tbl_id}}}"
                            )
                    elif tbl_obj and tbl_id in emitted_tables:
                        tbl_envs.append(
                            f"% Table~\\ref{{tab:{tbl_id}}} already defined above."
                        )
                    else:
                        tbl_envs.append(
                            f"% Table reference: {tbl_id} (no YAML — run paperforge add-table)"
                        )
                if tbl_envs:
                    text_par += "\n\n" + "\n\n".join(tbl_envs)

                alg_envs = []
                for alg_id in c.algorithms:
                    alg_obj = project.algorithm_map.get(alg_id)
                    if alg_obj and alg_id not in emitted_algorithms:
                        emitted_algorithms.add(alg_id)
                        alg_envs.append(alg_obj.to_latex())
                    elif alg_obj and alg_id in emitted_algorithms:
                        alg_envs.append(
                            f"% Algorithm~\\ref{{alg:{alg_id}}} already defined above."
                        )
                    else:
                        alg_envs.append(
                            f"% Algorithm reference: {alg_id} (no algorithm YAML)"
                        )
                if alg_envs:
                    text_par += "\n\n" + "\n\n".join(alg_envs)

                claim_blocks.append(text_par)

            if section == "introduction":
                if contrib:
                    items = []
                    for c in contrib:
                        emitted_claims.add(c.id)
                        items.append(f"  \\item {_claim_paragraph(c, project)}")
                    item_block = (
                        "\\noindent The main contributions of this work are:\n"
                        "\\begin{itemize}\n" + "\n".join(items) + "\n\\end{itemize}"
                    )
                    claim_blocks.append(item_block)

                if project.config.sections_overview:
                    so = _generate_sections_overview(project.config.sections_overview)
                    if so:
                        claim_blocks.append(so)

            block += "\n\n".join(claim_blocks)
        else:
            block += "% TODO: No claims linked to this section yet."
        blocks.append(block)
    return "\n\n".join(blocks)


def _ieee_parstart(text: str) -> str:
    if not text:
        return "\\IEEEPARstart{T}{his} section has no claims yet."
    first_word, _, rest_text = text.partition(" ")
    if len(first_word) < 2:
        drop_cap, rest_of_word = first_word, ""
    else:
        drop_cap, rest_of_word = first_word[0], first_word[1:]
    result = f"\\IEEEPARstart{{{drop_cap}}}{{{rest_of_word}}}"
    return f"{result} {rest_text}" if rest_text else result


def _generate_journal_sections(
    sections: list[str],
    project: PaperForgeProject,
    output_dir: Path | None = None,
    plugin: VenuePlugin | None = None,
) -> str:
    blocks: list[str] = []
    emitted_figures: set[str] = set()
    emitted_tables: set[str] = set()
    emitted_algorithms: set[str] = set()
    emitted_claims: set[str] = set()

    for section in sections:
        if section == "abstract":
            continue

        title = SECTION_TITLES.get(section, section.replace("_", " ").title())
        section_claims = sorted(
            (c for c in project.claims if section in c.sections), key=lambda c: c.id
        )
        current_subsection = ""

        if section == "introduction":
            policy = plugin.first_section_heading_policy if plugin else "raised_section"
            if policy == "raised_section":
                heading = (
                    f"\\IEEEraisesectionheading{{\\section{{{title}}}"
                    f"\\label{{sec:introduction}}}}"
                )
            else:
                heading = f"\\section{{{title}}}\\label{{sec:introduction}}"
            if section_claims:
                paragraphs: list[str] = []
                non_contrib = [c for c in section_claims if not c.is_contribution]
                contrib = [c for c in section_claims if c.is_contribution]
                target_claims = non_contrib if contrib else section_claims

                for idx, c in enumerate(target_claims):
                    if c.id in emitted_claims:
                        first_sec = c.sections[0] if c.sections else "above"
                        paragraphs.append(
                            f"% Claim {c.id} content appears in {first_sec} above."
                        )
                        continue
                    emitted_claims.add(c.id)

                    prefix = ""
                    if c.subsection and c.subsection != current_subsection:
                        prefix = f"\\subsection{{{escape_latex(c.subsection)}}}\n"
                        current_subsection = c.subsection

                    claim_text = _claim_paragraph(c, project)
                    if idx == 0 or not any(not p.startswith("%") for p in paragraphs):
                        first_text = prefix + _ieee_parstart(claim_text)
                    else:
                        first_text = prefix + claim_text

                    first_envs = []
                    for fig_id in c.figures:
                        fig_obj = next(
                            (f for f in project.figures if f.id == fig_id), None
                        )
                        if fig_obj and fig_id not in emitted_figures:
                            emitted_figures.add(fig_id)
                            first_envs.append(
                                _generate_figure_latex(
                                    fig_obj, project.root, output_dir
                                )
                            )
                        elif fig_obj and fig_id in emitted_figures:
                            first_envs.append(
                                f"% Figure~\\ref{{fig:{fig_id}}} already defined above."
                            )
                        else:
                            first_envs.append(
                                f"% Reference: {fig_id} (no figure YAML — run paperforge add-figure)"
                            )
                    if first_envs:
                        first_text += "\n\n" + "\n\n".join(first_envs)

                    first_tbl_envs = []
                    for tbl_id in c.tables:
                        tbl_obj = next(
                            (t for t in project.tables if t.id == tbl_id), None
                        )
                        if tbl_obj and tbl_id not in emitted_tables:
                            emitted_tables.add(tbl_id)
                            if tbl_obj.caption or tbl_obj.auto_rows_from_experiment:
                                first_tbl_envs.append(
                                    _generate_table_latex(tbl_obj, project)
                                )
                            else:
                                caption_text = escape_latex(
                                    (tbl_obj.caption or "")[:60]
                                )
                                first_tbl_envs.append(
                                    f"% Table: {tbl_id} — {caption_text} (no caption set)\n"
                                    f"% \\label{{tab:{tbl_id}}}"
                                )
                        elif tbl_obj and tbl_id in emitted_tables:
                            first_tbl_envs.append(
                                f"% Table~\\ref{{tab:{tbl_id}}} already defined above."
                            )
                        else:
                            first_tbl_envs.append(
                                f"% Table reference: {tbl_id} (no YAML — run paperforge add-table)"
                            )
                    if first_tbl_envs:
                        first_text += "\n\n" + "\n\n".join(first_tbl_envs)

                    alg_envs = []
                    for alg_id in c.algorithms:
                        alg_obj = project.algorithm_map.get(alg_id)
                        if alg_obj and alg_id not in emitted_algorithms:
                            emitted_algorithms.add(alg_id)
                            alg_envs.append(alg_obj.to_latex())
                        elif alg_obj and alg_id in emitted_algorithms:
                            alg_envs.append(
                                f"% Algorithm~\\ref{{alg:{alg_id}}} already defined above."
                            )
                        else:
                            alg_envs.append(
                                f"% Algorithm reference: {alg_id} (no algorithm YAML)"
                            )
                    if alg_envs:
                        first_text += "\n\n" + "\n\n".join(alg_envs)

                    paragraphs.append(first_text)

                if contrib:
                    items = []
                    for c in contrib:
                        emitted_claims.add(c.id)
                        items.append(f"  \\item {_claim_paragraph(c, project)}")
                    item_block = (
                        "\\noindent The main contributions of this work are:\n"
                        "\\begin{itemize}\n" + "\n".join(items) + "\n\\end{itemize}"
                    )
                    paragraphs.append(item_block)

                if project.config.sections_overview:
                    so = _generate_sections_overview(project.config.sections_overview)
                    if so:
                        paragraphs.append(so)

                body = "\n\n".join(paragraphs)
            else:
                body = "% TODO: No claims linked to this section yet."
            blocks.append(f"{heading}\n{body}")
            continue

        block = f"\\section{{{title}}}\n"
        if section_claims:
            claim_blocks = []
            for c in section_claims:
                if c.id in emitted_claims:
                    first_sec = c.sections[0] if c.sections else "above"
                    claim_blocks.append(
                        f"% Claim {c.id} content appears in {first_sec} above."
                    )
                    continue
                emitted_claims.add(c.id)

                prefix = ""
                if c.subsection and c.subsection != current_subsection:
                    prefix = f"\\subsection{{{escape_latex(c.subsection)}}}\n"
                    current_subsection = c.subsection

                if section == "related_work" and c.compared_work:
                    prefix += f"% --- {escape_latex(c.compared_work)} ---\n"

                text_par = prefix + _claim_paragraph(c, project)

                fig_envs = []
                for fig_id in c.figures:
                    fig_obj = next((f for f in project.figures if f.id == fig_id), None)
                    if fig_obj and fig_id not in emitted_figures:
                        emitted_figures.add(fig_id)
                        fig_envs.append(
                            _generate_figure_latex(fig_obj, project.root, output_dir)
                        )
                    elif fig_obj and fig_id in emitted_figures:
                        fig_envs.append(
                            f"% Figure~\\ref{{fig:{fig_id}}} already defined above."
                        )
                    else:
                        fig_envs.append(
                            f"% Reference: {fig_id} (no figure YAML — run paperforge add-figure)"
                        )
                if fig_envs:
                    text_par += "\n\n" + "\n\n".join(fig_envs)

                tbl_envs = []
                for tbl_id in c.tables:
                    tbl_obj = next((t for t in project.tables if t.id == tbl_id), None)
                    if tbl_obj and tbl_id not in emitted_tables:
                        emitted_tables.add(tbl_id)
                        if tbl_obj.caption or tbl_obj.auto_rows_from_experiment:
                            tbl_envs.append(_generate_table_latex(tbl_obj, project))
                        else:
                            caption_text = escape_latex((tbl_obj.caption or "")[:60])
                            tbl_envs.append(
                                f"% Table: {tbl_id} — {caption_text} (no caption set)\n"
                                f"% \\label{{tab:{tbl_id}}}"
                            )
                    elif tbl_obj and tbl_id in emitted_tables:
                        tbl_envs.append(
                            f"% Table~\\ref{{tab:{tbl_id}}} already defined above."
                        )
                    else:
                        tbl_envs.append(
                            f"% Table reference: {tbl_id} (no YAML — run paperforge add-table)"
                        )
                if tbl_envs:
                    text_par += "\n\n" + "\n\n".join(tbl_envs)

                alg_envs = []
                for alg_id in c.algorithms:
                    alg_obj = project.algorithm_map.get(alg_id)
                    if alg_obj and alg_id not in emitted_algorithms:
                        emitted_algorithms.add(alg_id)
                        alg_envs.append(alg_obj.to_latex())
                    elif alg_obj and alg_id in emitted_algorithms:
                        alg_envs.append(
                            f"% Algorithm~\\ref{{alg:{alg_id}}} already defined above."
                        )
                    else:
                        alg_envs.append(
                            f"% Algorithm reference: {alg_id} (no algorithm YAML)"
                        )
                if alg_envs:
                    text_par += "\n\n" + "\n\n".join(alg_envs)

                claim_blocks.append(text_par)
            block += "\n\n".join(claim_blocks)
        else:
            block += "% TODO: No claims linked to this section yet."
        blocks.append(block)
    return "\n\n".join(blocks)


def _get_membership_tag(membership: str | None) -> str:
    if not membership:
        return ""
    mem = membership.strip()
    valid_grades = {
        "Member",
        "Senior Member",
        "Fellow",
        "Life Member",
        "Life Fellow",
        "Life Senior Member",
    }
    if mem not in valid_grades:
        return ""
    return f",~\\IEEEmembership{{{mem},~IEEE}}"


def _generate_author_block_journal(
    authors: list[Any],
    affiliations: list[Affiliation],
    config: ProjectConfig | None = None,
    compsoc: bool = False,
) -> str:
    escaped_authors = [
        escape_latex(a.full_name if hasattr(a, "full_name") else str(a))
        for a in authors
    ]
    if not authors:
        return "Author(s) TBD"

    if compsoc:
        lines = []
        for i, author in enumerate(escaped_authors):
            aff = affiliations[i] if i < len(affiliations) else None
            author_obj = authors[i] if i < len(authors) else None
            mem_grade = getattr(author_obj, "ieee_membership_grade", None) or (
                aff.membership if aff else None
            )
            mem_tag = _get_membership_tag(mem_grade)
            if aff:
                aff_str = ", ".join(
                    filter(
                        None,
                        [
                            escape_latex(aff.department or ""),
                            escape_latex(aff.institution or ""),
                            escape_latex(aff.city or ""),
                            escape_latex(aff.country or ""),
                        ],
                    )
                )
                lines.append(f"  {author}{mem_tag}")
                if aff_str:
                    lines.append(
                        f"  \\IEEEcompsocitemizethanks{{"
                        f"\\IEEEcompsocthanksitem {author} is with "
                        f"{aff_str}.}}"
                    )
            else:
                lines.append(f"  {author}{mem_tag}")
        return "\n".join(lines)

    # Non-compsoc journal mode
    if not affiliations and (
        not config
        or (not config.funding and not config.email and not config.manuscript_received)
    ):
        return ", ".join(escaped_authors)

    has_extra = config and (
        config.funding or config.email or config.manuscript_received
    )

    if has_extra:
        author_parts = []
        for i, author in enumerate(escaped_authors):
            aff = affiliations[i] if i < len(affiliations) else None
            author_obj = authors[i] if i < len(authors) else None
            mem_grade = getattr(author_obj, "ieee_membership_grade", None) or (
                aff.membership if aff else None
            )
            mem_tag = _get_membership_tag(mem_grade)
            if config and config.orcid:
                author_parts.append(
                    f"{author}{mem_tag}~\\orcidlink{{{escape_latex(config.orcid)}}}"
                )
            else:
                author_parts.append(f"{author}{mem_tag}")

        if len(author_parts) == 1:
            author_str = author_parts[0]
        elif len(author_parts) == 2:
            author_str = f"{author_parts[0]} and {author_parts[1]}"
        else:
            author_str = ", ".join(author_parts[:-1]) + f", and {author_parts[-1]}"

        thanks_parts = []
        if config and config.funding:
            fund_text = escape_latex(config.funding).strip()
            if not fund_text.endswith("."):
                fund_text += "."
            thanks_parts.append(fund_text)

        if affiliations:
            aff_lines = []
            for i, aff in enumerate(affiliations):
                aff_items = [
                    escape_latex(aff.department or ""),
                    escape_latex(aff.institution or ""),
                    escape_latex(aff.city or ""),
                    escape_latex(aff.country or ""),
                ]
                aff_str = ", ".join(filter(None, aff_items))
                if aff.email:
                    aff_str += f" (e-mail: {escape_latex(aff.email)})"
                if aff_str:
                    if len(escaped_authors) > i:
                        aff_lines.append(f"{escaped_authors[i]} is with {aff_str}.")
                    else:
                        aff_lines.append(f"{aff_str}.")
            if aff_lines:
                thanks_parts.append(" ".join(aff_lines))

        if config and config.email:
            thanks_parts.append(
                f"Corresponding author: {escaped_authors[0]} (e-mail: {escape_latex(config.email)})."
            )

        if config and config.manuscript_received:
            thanks_parts.append(
                f"Manuscript received {escape_latex(config.manuscript_received)}."
            )

        thanks_block = " ".join(thanks_parts)
        return f"{author_str}\\thanks{{{thanks_block}}}"
    else:
        parts = []
        for i, author in enumerate(escaped_authors):
            aff = affiliations[i] if i < len(affiliations) else None
            author_obj = authors[i] if i < len(authors) else None
            mem_grade = getattr(author_obj, "ieee_membership_grade", None) or (
                aff.membership if aff else None
            )
            mem_tag = _get_membership_tag(mem_grade)
            if config and config.orcid:
                author_fmt = f"{author}~\\orcidlink{{{escape_latex(config.orcid)}}}"
            else:
                author_fmt = author
            if aff:
                aff_items = [
                    escape_latex(aff.department or ""),
                    escape_latex(aff.institution or ""),
                    escape_latex(aff.city or ""),
                    escape_latex(aff.country or ""),
                ]
                aff_str = ", ".join(filter(None, aff_items))
                if aff.email:
                    aff_str += f" (e-mail: {escape_latex(aff.email)})"
                if aff_str:
                    parts.append(f"{author_fmt}{mem_tag}\\thanks{{{aff_str}}}")
                else:
                    parts.append(f"{author_fmt}{mem_tag}")
            else:
                parts.append(f"{author_fmt}{mem_tag}")
        return " \\and\n".join(parts)


def _generate_acknowledgment(project: PaperForgeProject) -> str:
    ack_text = project.config.acknowledgment
    comment_line = ""
    if project.config.funding:
        comment_line = (
            "% Note: Funding acknowledgment is in the \\thanks{} footnote per IEEE convention. "
            "Add only people/institution thanks here.\n"
        )
    if not ack_text:
        ack_text = "% TODO: Add acknowledgment text."
    else:
        ack_text = escape_latex(ack_text)
    return (
        f"{comment_line}"
        "\\ifCLASSOPTIONcompsoc\n"
        "  \\section*{Acknowledgments}\n"
        "\\else\n"
        "  \\section*{Acknowledgment}\n"
        "\\fi\n\n"
        f"{ack_text}"
    )


def _generate_bibliography_from_citations(
    project: PaperForgeProject,
    all_keys: set[str],
) -> str:
    lines = [
        "% references.bib — generated by PaperForge",
        "% Source: .paperforge/citations/*.yaml",
        "% Re-generated on every build when citation YAMLs exist.",
        "",
    ]
    cit_map = project.citation_map
    for key in sorted(all_keys):
        if key in cit_map:
            lines.append(cit_map[key].to_bibtex())
        else:
            lines.append(
                f"@article{{{key},\n"
                f"  author  = {{Author, A.}},\n"
                f"  title   = {{TODO: fill in title for {key}}},\n"
                f"  journal = {{TODO}},\n"
                f"  year    = {{2024}},\n"
                f"  note    = {{Add .paperforge/citations/{key}.yaml for real metadata}},\n"
                f"}}"
            )
        lines.append("")
    return "\n".join(lines)


def _generate_bibliography_stubs(all_keys: set[str]) -> str:
    entries = [
        f"@article{{{key},\n"
        f"  author = {{Author, A.}},\n"
        f"  title  = {{TODO: Title for {key}}},\n"
        f"  journal = {{TODO}},\n"
        f"  year   = {{2024}},\n"
        f"  note   = {{Auto-generated stub. Replace with real entry.}},\n"
        f"}}"
        for key in sorted(all_keys)
    ]
    return "\n\n".join(entries) + "\n"


def _generate_bibliography(project: PaperForgeProject) -> str:
    """Return bibliography LaTeX block."""
    unique_citations = sorted(
        {citation for claim in project.claims for citation in claim.citations}
    )
    if unique_citations or project.citations:
        return "\\bibliographystyle{IEEEtran}\n\\bibliography{references}"

    return (
        "\\begin{thebibliography}{99}\n"
        "\\bibitem{ref1} Placeholder reference. Replace with a real citation.\n"
        "\\end{thebibliography}"
    )


def _bib_has_real_entries(bib_path: Path) -> bool:
    """
    Returns True if the .bib file contains at least one
    real entry (an @-block that does not contain 'TODO').
    Returns False if file doesn't exist or only has stubs.
    """
    if not bib_path.exists():
        return False
    content = bib_path.read_text(encoding="utf-8")
    entries = re.findall(r"@\w+\{[^}]+\}", content, re.DOTALL)
    for entry in entries:
        if "TODO" not in entry:
            return True
    return False


def _generate_biographies(biographies: list[Biography]) -> str:
    """Return the biography block LaTeX or empty string."""
    if not biographies:
        return ""
    parts = [b.to_latex() for b in biographies]
    return "\n\n".join(parts)


def _generate_ai_disclosure(text: str) -> str:
    """Return the AI disclosure section or empty string."""
    if not text:
        return ""
    from paperforge.utils.latex import escape_latex_safe

    return "\\subsection*{Use of Artificial Intelligence Tools}\n" + escape_latex_safe(
        text
    )


def _generate_latex_conference(
    project: PaperForgeProject, plugin: VenuePlugin, output_dir: Path | None = None
) -> str:
    """Generate LaTeX for a conference-style IEEE/ACM/NeurIPS paper."""
    title = escape_latex(project.config.title or "Untitled Paper")
    escaped_authors = [
        escape_latex(a.full_name if hasattr(a, "full_name") else str(a))
        for a in project.config.authors
    ]
    author_block = plugin.generate_author_block(escaped_authors)
    abstract_content = _generate_abstract(project.claims)
    sections = _generate_sections(project.config.sections, project, output_dir)
    bibliography = _generate_bibliography(project)

    preamble = plugin.generate_preamble()
    if "\\usepackage{algorithmic}" not in preamble:
        preamble += "\n\\usepackage{algorithm}\n\\usepackage{algorithmic}"
    if project.config.orcid and "\\usepackage{orcidlink}" not in preamble:
        preamble += "\n\\usepackage{orcidlink}"
    if project.config.theorem_packages and "\\newtheorem" not in preamble:
        preamble += (
            "\n\\usepackage{amsthm}\n"
            "\\newtheorem{theorem}{Theorem}[section]\n"
            "\\newtheorem{lemma}[theorem]{Lemma}\n"
            "\\newtheorem{definition}[theorem]{Definition}\n"
            "\\newtheorem{corollary}[theorem]{Corollary}\n"
            "\\newtheorem{remark}{Remark}\n"
            "\\theoremstyle{remark}"
        )

    first_author = project.config.authors[0] if project.config.authors else None
    author_str = (
        first_author.cite_name
        if first_author and hasattr(first_author, "cite_name")
        else (str(first_author) if first_author else "")
    )
    hypersetup_block = f"""\\hypersetup{{
    pdftitle={{{escape_latex(project.config.title)}}},
    pdfauthor={{{escape_latex(author_str)}}},
    pdfsubject={{{escape_latex(project.config.venue)}}},
    pdfkeywords={{{escape_latex(", ".join(project.config.keywords))}}},
    hidelinks
}}"""
    preamble += "\n" + hypersetup_block

    statements = []
    if project.config.data_availability:
        statements.append(
            f"\\section*{{Data Availability}}\n{escape_latex(project.config.data_availability)}"
        )
    if project.config.code_availability:
        statements.append(
            f"\\section*{{Code Availability}}\n{escape_latex(project.config.code_availability)}"
        )
    if project.config.conflict_of_interest:
        statements.append(
            f"\\section*{{Conflict of Interest}}\n{escape_latex(project.config.conflict_of_interest)}"
        )
    statements_block = "\n\n".join(statements)
    if statements_block:
        statements_block = "\n\n" + statements_block

    ai_disclosure = _generate_ai_disclosure(project.config.ai_disclosure)
    if ai_disclosure:
        statements_block += "\n\n" + ai_disclosure

    bio_block = _generate_biographies(project.config.biographies)
    if bio_block:
        bio_block = "\n\n\\vfill\n\\newpage\n\n" + bio_block

    return f"""{plugin.latex_documentclass}

{preamble}

\\begin{{document}}

\\title{{{title}}}

{author_block}

\\maketitle

\\begin{{abstract}}
{abstract_content}
\\end{{abstract}}

{sections}{statements_block}

{bibliography}{bio_block}

\\end{{document}}
"""


def _generate_latex_journal(
    project: PaperForgeProject, plugin: VenuePlugin, output_dir: Path | None = None
) -> str:
    """Generate LaTeX for an IEEE Transactions / journal paper."""
    title = escape_latex(project.config.title or "Untitled Paper")
    compsoc = getattr(plugin, "mode", None) == "journal-compsoc"
    author_block = _generate_author_block_journal(
        project.config.authors,
        project.config.affiliations,
        project.config,
        compsoc=compsoc,
    )
    abstract_content = _generate_abstract(project.claims)
    keywords_source = project.config.keywords or project.config.sections[:6]
    keywords = ", ".join(escape_latex(k) for k in keywords_source)
    sections = _generate_journal_sections(
        project.config.sections, project, output_dir, plugin
    )
    bibliography = _generate_bibliography(project)
    acknowledgment = _generate_acknowledgment(project)

    preamble = plugin.generate_preamble()
    if "\\usepackage{algorithmic}" not in preamble:
        preamble += "\n\\usepackage{algorithm}\n\\usepackage{algorithmic}"
    if project.config.orcid and "\\usepackage{orcidlink}" not in preamble:
        preamble += "\n\\usepackage{orcidlink}"
    if project.config.theorem_packages and "\\newtheorem" not in preamble:
        preamble += (
            "\n\\usepackage{amsthm}\n"
            "\\newtheorem{theorem}{Theorem}[section]\n"
            "\\newtheorem{lemma}[theorem]{Lemma}\n"
            "\\newtheorem{definition}[theorem]{Definition}\n"
            "\\newtheorem{corollary}[theorem]{Corollary}\n"
            "\\newtheorem{remark}{Remark}\n"
            "\\theoremstyle{remark}"
        )

    first_author = project.config.authors[0] if project.config.authors else None
    author_str = (
        first_author.cite_name
        if first_author and hasattr(first_author, "cite_name")
        else (str(first_author) if first_author else "")
    )
    hypersetup_block = f"""\\hypersetup{{
    pdftitle={{{escape_latex(project.config.title)}}},
    pdfauthor={{{escape_latex(author_str)}}},
    pdfsubject={{{escape_latex(project.config.venue)}}},
    pdfkeywords={{{escape_latex(", ".join(project.config.keywords))}}},
    hidelinks
}}"""
    preamble += "\n" + hypersetup_block

    publisher_id_block = ""
    if project.config.publisher_id:
        publisher_id_block = (
            f"\n\\IEEEpubid{{{escape_latex(project.config.publisher_id)}}}\n"
        )

    statements = []
    if project.config.data_availability:
        statements.append(
            f"\\section*{{Data Availability}}\n{escape_latex(project.config.data_availability)}"
        )
    if project.config.code_availability:
        statements.append(
            f"\\section*{{Code Availability}}\n{escape_latex(project.config.code_availability)}"
        )
    if project.config.conflict_of_interest:
        statements.append(
            f"\\section*{{Conflict of Interest}}\n{escape_latex(project.config.conflict_of_interest)}"
        )
    statements_block = "\n\n".join(statements)
    if statements_block:
        statements_block = "\n\n" + statements_block

    ai_disclosure = _generate_ai_disclosure(project.config.ai_disclosure)
    if ai_disclosure:
        statements_block += "\n\n" + ai_disclosure

    bio_block = _generate_biographies(project.config.biographies)
    if bio_block:
        bio_block = "\n\n\\vfill\n\\newpage\n\n" + bio_block

    is_ieee_access = (project.config.venue or "").lower() == "ieee access"
    journal_id_cmd = "\\journalid{IEEE Access}" if is_ieee_access else ""
    kw_env = "keywords" if is_ieee_access else "IEEEkeywords"

    return f"""{plugin.latex_documentclass}

{preamble}

\\providecommand{{\\journalid}}[1]{{}}
\\newenvironment{{keywords}}{{\\begin{{IEEEkeywords}}}}{{\\end{{IEEEkeywords}}}}

\\hyphenation{{op-tical net-works semi-conduc-tor}}

\\begin{{document}}
{journal_id_cmd}

\\title{{{title}}}

\\author{{{author_block}}}

\\IEEEtitleabstractindextext{{%
\\begin{{abstract}}
{abstract_content}
\\end{{abstract}}

\\begin{{{kw_env}}}
{keywords}
\\end{{{kw_env}}}}}

\\maketitle
{publisher_id_block}
\\IEEEdisplaynontitleabstractindextext
\\IEEEpeerreviewmaketitle

{sections}

{acknowledgment}{statements_block}

{bibliography}{bio_block}

\\end{{document}}
"""


def _is_pdf_stale(project_root: Path, output_dir: Path | None = None) -> bool:
    """
    Returns True if paper.pdf does not exist or is older
    than any source file in .paperforge/.
    Returns False if PDF is newer than all sources.
    """
    if output_dir is None:
        output_dir = project_root / "paper_generated" / "current"
        if (
            not (output_dir / "paper.pdf").exists()
            and (project_root / "paper" / "paper.pdf").exists()
        ):
            output_dir = project_root / "paper"
    pdf_path = output_dir / "paper.pdf"
    if not pdf_path.exists():
        return True

    pdf_mtime = pdf_path.stat().st_mtime
    pf_dir = project_root / ".paperforge"

    # Check all YAML source files
    for yaml_file in pf_dir.rglob("*.yaml"):
        if yaml_file.stat().st_mtime > pdf_mtime:
            return True

    return False


def _compile_pdf_full(
    tex_path: Path,
    output_dir: Path,
) -> tuple[bool, str]:
    latexmk = shutil.which("latexmk")
    pdflatex = shutil.which("pdflatex")
    bibtex = shutil.which("bibtex")

    if latexmk:
        # latexmk handles the full BibTeX pipeline automatically
        result = subprocess.run(
            [
                latexmk,
                "-pdf",
                "-bibtex",
                "-interaction=nonstopmode",
                f"-outdir={output_dir}",
                str(tex_path),
            ],
            capture_output=True,
            text=True,
            cwd=output_dir,
            check=False,
        )
        return result.returncode == 0, "latexmk"

    if pdflatex:
        tex_name = tex_path.stem  # "paper" without extension

        def run_pdflatex() -> subprocess.CompletedProcess[str]:
            return subprocess.run(
                [
                    pdflatex,
                    "-interaction=nonstopmode",
                    f"-output-directory={output_dir}",
                    str(tex_path),
                ],
                capture_output=True,
                text=True,
                cwd=output_dir,
                check=False,
            )

        def run_bibtex() -> subprocess.CompletedProcess[str] | None:
            if bibtex:
                return subprocess.run(
                    [bibtex, tex_name],
                    capture_output=True,
                    text=True,
                    cwd=output_dir,
                    check=False,
                )
            return None

        # Full BibTeX pipeline: pdflatex -> bibtex -> pdflatex -> pdflatex
        run_pdflatex()  # pass 1: generate .aux
        run_bibtex()  # bibtex: process references
        run_pdflatex()  # pass 2: resolve citations
        result = run_pdflatex()  # pass 3: resolve cross-references
        return result.returncode == 0, "pdflatex+bibtex"

    return False, "none"


_compile_pdf = _compile_pdf_full


def _generate_docx(
    project: PaperForgeProject,
    output_dir: Path,
) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()

    # Page margins (IEEE-like: narrow)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title_para.add_run(project.config.title or "Untitled")
    run_t.bold = True
    run_t.font.size = Pt(14)

    # Authors
    authors_para = doc.add_paragraph()
    authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_para.add_run(
        ", ".join(str(a) for a in project.config.authors) or "Author TBD"
    )

    # Affiliations
    for aff in project.config.affiliations:
        aff_parts = list(
            filter(
                None,
                [aff.department, aff.institution, aff.city, aff.country],
            )
        )
        aff_str = ", ".join(aff_parts)
        if aff_str:
            aff_para = doc.add_paragraph()
            aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_a = aff_para.add_run(aff_str)
            run_a.italic = True
            run_a.font.size = Pt(9)

    doc.add_paragraph()  # spacing

    # Abstract
    abs_claims = [c for c in project.claims if "abstract" in c.sections]
    if abs_claims:
        abs_para = doc.add_paragraph()
        abs_run = abs_para.add_run("Abstract\u2014")
        abs_run.bold = True
        abs_para.add_run(" ".join(c.text for c in abs_claims if c.text))

    # Keywords
    if project.config.keywords:
        kw_para = doc.add_paragraph()
        kw_run = kw_para.add_run("Index Terms\u2014")
        kw_run.bold = True
        kw_para.add_run(", ".join(project.config.keywords))

    doc.add_paragraph()

    # Sections
    section_titles = {
        "introduction": "Introduction",
        "related_work": "Related Work",
        "methodology": "Methodology",
        "experiments": "Experimental Setup",
        "results": "Results",
        "discussion": "Discussion",
        "conclusion": "Conclusion",
    }

    emitted_claims: set[str] = set()
    emitted_tables: set[str] = set()

    for section_name in project.config.sections:
        if section_name == "abstract":
            continue

        title = section_titles.get(section_name, section_name.replace("_", " ").title())
        doc.add_heading(title, level=1)

        section_claims = [c for c in project.claims if section_name in c.sections]

        for claim in section_claims:
            if claim.id in emitted_claims:
                continue
            if claim.text:
                doc.add_paragraph(claim.text)
                emitted_claims.add(claim.id)

            # Tables
            for tbl_id in claim.tables:
                if tbl_id in emitted_tables:
                    continue
                tbl_obj = next((t for t in project.tables if t.id == tbl_id), None)
                if tbl_obj and tbl_obj.columns and tbl_obj.rows:
                    # Caption above table (IEEE style)
                    cap = doc.add_paragraph()
                    cap.add_run(f"TABLE: {tbl_obj.caption}").bold = True
                    # Table
                    word_table = doc.add_table(
                        rows=1 + len(tbl_obj.rows), cols=len(tbl_obj.columns)
                    )
                    word_table.style = "Table Grid"
                    # Header row
                    hdr = word_table.rows[0].cells
                    for i, col in enumerate(tbl_obj.columns):
                        hdr[i].text = col
                        hdr[i].paragraphs[0].runs[0].bold = True
                    # Data rows
                    for r_idx, row in enumerate(tbl_obj.rows):
                        cells = word_table.rows[r_idx + 1].cells
                        for c_idx, cell in enumerate(row):
                            if c_idx < len(cells):
                                cells[c_idx].text = cell
                    if tbl_obj.notes:
                        note_para = doc.add_paragraph()
                        run_n = note_para.add_run(tbl_obj.notes)
                        run_n.italic = True
                        run_n.font.size = Pt(8)
                    emitted_tables.add(tbl_id)

    # Acknowledgment
    ack = project.config.acknowledgment
    if ack and ack.strip():
        doc.add_heading("Acknowledgment", level=1)
        doc.add_paragraph(ack)

    # COI
    coi = project.config.conflict_of_interest
    if coi and coi.strip():
        doc.add_heading("Conflict of Interest", level=1)
        doc.add_paragraph(coi)

    # References
    doc.add_heading("References", level=1)
    cit_map = project.citation_map
    all_keys = sorted({key for claim in project.claims for key in claim.citations})
    for i, key in enumerate(all_keys, 1):
        cit = cit_map.get(key)
        if cit and cit.title:
            authors_str = " and ".join(cit.authors) if cit.authors else "Author"
            ref_text = (
                f'[{i}] {authors_str}, "{cit.title}," {cit.venue}, {cit.year or "n.d."}'
            )
            if cit.doi:
                ref_text += f", doi: {cit.doi}"
        else:
            ref_text = f"[{i}] {key} — add citation YAML for real reference"
        doc.add_paragraph(ref_text)

    docx_path = output_dir / "paper.docx"
    doc.save(str(docx_path))
    return docx_path


def _reveal_output(path: Path) -> None:
    """Open the containing folder of the output file."""
    try:
        system = platform.system()
        if system == "Windows":
            subprocess.run(["explorer", "/select,", str(path)], check=False)
        elif system == "Darwin":
            subprocess.run(["open", "-R", str(path)], check=False)
        elif system == "Linux":
            subprocess.run(["xdg-open", str(path.parent)], check=False)
    except (OSError, subprocess.SubprocessError):
        pass  # Never crash the build over a reveal failure


def _cleanup_aux_files(output_dir: Path) -> None:
    """Remove LaTeX auxiliary files after compilation."""
    aux_extensions = {
        ".aux",
        ".log",
        ".fls",
        ".fdb_latexmk",
        ".out",
        ".bbl",
        ".blg",
        ".synctex.gz",
        ".toc",
        ".lof",
        ".lot",
    }
    for f in output_dir.iterdir():
        if f.is_file() and f.suffix in aux_extensions:
            try:
                f.unlink()
            except OSError:
                pass


def _path_contains(parent: Path, candidate: Path) -> bool:
    """True if candidate is parent itself or nested inside parent."""
    return candidate == parent or parent in candidate.parents


def _rotate_output(
    project_root: Path,
    output_dir: Path | None = None,
    policy: str | None = None,
    archive_dir: Path | None = None,
) -> None:
    """Archive the previous build of a single output directory before a new build.

    Rotation is scoped to the selected `output_dir` only: it never touches
    any other output directory (e.g. a sibling candidate build directory),
    and it never guesses a shared archive location across different output
    directory names. Only copies meaningful files -- never aux files.
    """
    project: PaperForgeProject | None = None
    if output_dir is None:
        project = PaperForgeProject.load(project_root)
        rel_output = project.config.build_output_dir or "paper_generated/current"
        output_dir = project_root / rel_output

    if policy is None and archive_dir is None:
        # Only consult project config when neither policy nor an explicit
        # archive directory was given -- an explicit archive_dir is a
        # direct instruction that should not require a project on disk.
        if project is None:
            project = PaperForgeProject.load(project_root)
        policy = project.config.output_rotation or "preserve_previous"
        if project.config.output_rotation_archive_dir:
            archive_dir = project_root / project.config.output_rotation_archive_dir
    elif policy is None:
        policy = "preserve_previous"

    if policy == "disabled":
        return

    output_dir.mkdir(parents=True, exist_ok=True)

    if archive_dir is not None:
        previous_dir = archive_dir
    elif policy == "timestamped":
        from datetime import datetime

        stamp = datetime.now(UTC).strftime("%Y%m%d_%H%M%S")
        previous_dir = output_dir.parent / f"{output_dir.name}_archive_{stamp}"
    else:
        # "preserve_previous" (default): scope the archive to this output
        # directory's own name, so sibling output directories (e.g. two
        # different candidate build targets sharing a parent) never share
        # -- and therefore never clobber -- each other's archive. The
        # conventional "current" -> "previous" pairing is kept for
        # backward compatibility.
        archive_name = (
            "previous"
            if output_dir.name == "current"
            else f"{output_dir.name}.previous"
        )
        previous_dir = output_dir.parent / archive_name

    try:
        output_resolved = output_dir.resolve()
        previous_resolved = previous_dir.resolve()
    except OSError:
        return

    # Reject unsafe rotation targets: archive nested inside output (or vice
    # versa), or archive equal to output.
    if _path_contains(output_resolved, previous_resolved) or _path_contains(
        previous_resolved, output_resolved
    ):
        return

    previous_dir.mkdir(parents=True, exist_ok=True)

    keep_files = [
        "paper.pdf",
        "paper_overleaf.zip",
        "paper.tex",
        "references.bib",
        "paper.docx",
        "traceability.tex",
    ]

    for filename in keep_files:
        src = output_dir / filename
        dst = previous_dir / filename
        if src.exists():
            shutil.copy2(str(src), str(dst))


DRAFT_BLOCKING = {
    "ORPHAN_CLAIM",
    "EMPTY_CLAIM_TEXT",
    "RESULTS_SECTION_EMPTY",
    "TABLE_NO_CAPTION",
    "CITATION_NO_TITLE",
    "LATEX_ARTIFACT_IN_CLAIM",
    "REQUIRED_PLACEHOLDER_IN_CLAIM",
    "AUTHOR_NAME_INCOMPLETE",
}

SUBMISSION_BLOCKING = DRAFT_BLOCKING | {
    "METRIC_CLAIM_MISMATCH",
    "ABSTRACT_HAS_CITATION",
    "ABSTRACT_INTRO_OVERLAP",
    "DUPLICATE_CLAIM_TEXT",
    "AUTHOR_IDENTITY_INCONSISTENT",
    "CITATION_HAS_INTERNAL_NOTE",
    "CLAIM_CONSTRAINT_VIOLATED",
    "PVALUE_AMBIGUOUS",
}


def _check_latex_artifacts(content: str) -> list[str]:
    """Scan generated LaTeX for artifacts that should not appear in final output."""
    issues = []
    patterns = [
        (r"\*\*\w[^*]*\*\*", "unresolved **bold** markdown"),
        (r"`[^`\n]+`", "unresolved `code` markdown"),
        (r"\[[\w\s]+\]\(https?://", "unresolved [link](url) markdown"),
        (r"(?<!\{)TODO(?!\})", "TODO placeholder in output"),
        (r"D-\d{3}", "internal decision ID in output"),
    ]
    lines = content.split("\n")
    for i, line in enumerate(lines, 1):
        if line.strip().startswith("%"):
            continue  # skip LaTeX comments
        for pattern, desc in patterns:
            if re.search(pattern, line):
                issues.append(f"Line {i}: {desc}: {line.strip()[:60]}")
    return issues


def _generate_build_reports(
    project: PaperForgeProject,
    issues: list[Any],
    output_dir: Path,
    mode: str,
) -> None:
    reports_dir = (
        output_dir.parent.parent / "reports"
        if output_dir.parent.name == "paper_generated"
        else output_dir / "reports"
    )
    reports_dir.mkdir(parents=True, exist_ok=True)

    from datetime import datetime

    ts = datetime.now(tz=UTC).strftime("%Y-%m-%d %H:%M")

    # doctor.md & doctor.json
    errors = [i for i in issues if getattr(i, "severity", "") == "ERROR"]
    warnings = [i for i in issues if getattr(i, "severity", "") == "WARNING"]
    infos = [i for i in issues if getattr(i, "severity", "") == "INFO"]

    doc_lines = [
        f"# Doctor Report — {ts}\n\n",
        f"Mode: {mode}\n",
        f"Total issues: {len(issues)}\n\n",
        f"## Errors ({len(errors)})\n\n",
        "| Code | Message |\n|------|---------|\n",
    ]
    for i in errors:
        doc_lines.append(f"| {i.code} | {i.message} |\n")

    doc_lines.append(
        f"\n## Warnings ({len(warnings)})\n\n| Code | Message |\n|------|---------|\n"
    )
    for i in warnings:
        doc_lines.append(f"| {i.code} | {i.message} |\n")

    doc_lines.append(
        f"\n## Info ({len(infos)})\n\n| Code | Message |\n|------|---------|\n"
    )
    for i in infos:
        doc_lines.append(f"| {i.code} | {i.message} |\n")

    (reports_dir / "doctor.md").write_text("".join(doc_lines), encoding="utf-8")
    doctor_json_data = [
        {
            "code": getattr(i, "code", ""),
            "severity": getattr(i, "severity", ""),
            "message": getattr(i, "message", ""),
            "claim_id": getattr(i, "claim_id", ""),
        }
        for i in issues
    ]
    (reports_dir / "doctor.json").write_text(
        json.dumps(doctor_json_data, indent=2), encoding="utf-8"
    )

    # Run services to produce remaining reports
    from paperforge.services.pdf_preflight import run_pdf_preflight
    from paperforge.services.reference_verifier import verify_references
    from paperforge.services.structural_integrity import check_structural_integrity
    from paperforge.services.template_fingerprint import verify_template_fingerprint

    tex_file = output_dir / "paper.tex"
    tex_text = tex_file.read_text(encoding="utf-8") if tex_file.exists() else ""
    fp_res = verify_template_fingerprint(
        tex_text, project.config.venue or "ieee", output_dir
    )
    (reports_dir / "venue_fingerprint.json").write_text(
        json.dumps(fp_res.to_dict(), indent=2), encoding="utf-8"
    )
    (reports_dir / "venue_fingerprint.md").write_text(
        f"# Venue Fingerprint Report — {ts}\n\n- Status: {fp_res.status}\n- Venue: {fp_res.requested_venue}\n- Detected: {fp_res.detected_template}\n",
        encoding="utf-8",
    )

    check_structural_integrity(project, reports_dir, mode=mode, tex_content=tex_text)
    verify_references(project, reports_dir, online=False)

    pdf_file = output_dir / "paper.pdf"
    if pdf_file.exists():
        run_pdf_preflight(pdf_file, reports_dir, mode=mode)

    # claim_evidence_report.md
    ev_lines = [
        f"# Claim Evidence Report — {ts}\n\n",
        "| Claim ID | Text (60 chars) | Experiment | Verified |\n",
        "|----------|-----------------|------------|---------|\n",
    ]
    for c in sorted(project.claims, key=lambda claim: claim.id):
        txt = (c.text[:57] + "...") if len(c.text) > 60 else c.text
        verified = "Yes" if c.status == "verified" else "No"
        exp = c.experiment or ", ".join(c.experiments) or "None"
        ev_lines.append(f"| {c.id} | {txt} | {exp} | {verified} |\n")

    (reports_dir / "claim_evidence_report.md").write_text(
        "".join(ev_lines), encoding="utf-8"
    )

    # submission_checklist.md
    issue_codes = {getattr(i, "code", "") for i in issues}
    chk_lines = [
        f"# Submission Checklist — {ts}\n\n",
        "## Critical (must pass for submission mode)\n\n",
    ]
    for code in sorted(SUBMISSION_BLOCKING):
        status = "[ ]" if code in issue_codes else "[x]"
        chk_lines.append(f"- {status} {code}\n")

    chk_lines.append("\n## Warnings (review before submission)\n\n")
    for i in warnings:
        chk_lines.append(f"- [ ] {i.code}: {i.message}\n")
    if not warnings:
        chk_lines.append("- [x] No warnings found\n")

    (reports_dir / "submission_checklist.md").write_text(
        "".join(chk_lines), encoding="utf-8"
    )


def run(
    project_root: Path,
    target: str = "ieee",
    no_reveal: bool = False,
    force: bool = False,
    force_anyway: bool = False,
    mode: str = "draft",
) -> None:
    if not (project_root / ".paperforge").exists():
        console.print(
            "[red]Not a PaperForge project. Run `paperforge init` first.[/red]"
        )
        sys.exit(1)

    project = PaperForgeProject.load(project_root)

    try:
        plugin = get_plugin(target)
    except ValueError as exc:
        console.print(f"[red]{exc}[/red]")
        sys.exit(1)

    issues = collect_issues(project, mode=mode)
    venue_issues = plugin.validate(project)

    submission_mode = mode == "submission"
    blocking_codes = SUBMISSION_BLOCKING if submission_mode else DRAFT_BLOCKING
    blocking_issues = [issue for issue in issues if issue.code in blocking_codes]
    venue_errors = [issue for issue in venue_issues if issue.severity == "ERROR"]
    if (blocking_issues or venue_errors) and not force_anyway:
        body = Group(
            Text(
                f"Build blocked ({mode} mode). Fix all ERRORs and blocking checks before building."
            ),
            *(Text(f"  [{issue.code}] {issue.message}") for issue in blocking_issues),
            *(Text(f"  [{issue.code}] {issue.message}") for issue in venue_errors),
            Text("Run `paperforge doctor` for full details."),
            Text("Use --force-anyway to bypass (NOT recommended for submission)."),
        )
        console.print(Panel(body, border_style="red"))
        sys.exit(1)

    rel_output = project.config.build_output_dir or "paper/paper_generated/current"
    output_dir = project_root / rel_output

    stale_root = project_root / "paper_generated"
    try:
        if stale_root.exists() and not output_dir.is_relative_to(stale_root):
            console.print(
                "[yellow]Found stale paper_generated/ at project root. "
                "Run `paperforge clean` to remove it.[/yellow]"
            )
    except (AttributeError, TypeError):
        pass

    _rotate_output(project_root, output_dir)
    pdf_path = output_dir / "paper.pdf"

    stale = _is_pdf_stale(project_root, output_dir)

    if not stale and pdf_path.exists() and not force:
        body = Group(
            Text(f"{rel_output}/paper.pdf is newer than all source files."),
            Text("No rebuild needed."),
            Text(""),
            Text("To force a rebuild: paperforge build --force"),
        )
        console.print(Panel(body, title="PDF Up To Date", border_style="dim"))
        return

    if stale and pdf_path.exists():
        console.print("[dim]Source changed — deleting stale PDF...[/dim]")
        pdf_path.unlink()

    venue_warnings = [issue for issue in venue_issues if issue.severity == "WARNING"]
    output_dir.mkdir(parents=True, exist_ok=True)

    unresolved_figures = [
        fig_obj
        for fig_obj in project.figures
        if fig_obj.path and resolve_figure_asset(fig_obj, project_root) is None
    ]
    if unresolved_figures:
        if submission_mode and not force_anyway:
            body = Group(
                Text(
                    "Build blocked (submission mode): figure asset(s) could not be resolved."
                ),
                *(
                    Text(
                        f"  [{fig_obj.id}] configured path: '{fig_obj.path}' "
                        f"-> not found under project root: {project_root}"
                    )
                    for fig_obj in unresolved_figures
                ),
                Text(
                    "Fix the figure path(s) in .paperforge/figures/*.yaml, or add the missing asset file."
                ),
                Text("Use --force-anyway to bypass (NOT recommended for submission)."),
            )
            console.print(Panel(body, border_style="red"))
            sys.exit(1)
        for fig_obj in unresolved_figures:
            console.print(
                f"[yellow]Warning: figure '{fig_obj.id}' asset not found "
                f"(configured path: '{fig_obj.path}', resolved against project root "
                f"'{project_root}'). A placeholder will be emitted.[/yellow]"
            )

    if project.config.paper_type == "journal":
        latex = _generate_latex_journal(project, plugin, output_dir)
    else:
        latex = _generate_latex_conference(project, plugin, output_dir)
    tex_path = output_dir / "paper.tex"
    tex_path.write_text(latex, encoding="utf-8")

    artifacts = _check_latex_artifacts(latex)
    if artifacts and not force_anyway:
        console.print("[red]LaTeX artifact check failed:[/red]")
        for a in artifacts[:10]:
            console.print(f"  {a}")
        console.print("Fix these before building. Use --force-anyway to override.")
        sys.exit(1)

    all_claim_citation_keys = {
        key for claim in project.claims for key in claim.citations
    }
    has_citation_yamls = len(project.citations) > 0
    bib_path = output_dir / "references.bib"

    if has_citation_yamls:
        bib_content = _generate_bibliography_from_citations(
            project, all_claim_citation_keys
        )
        bib_path.write_text(bib_content, encoding="utf-8")
    elif all_claim_citation_keys:
        if not _bib_has_real_entries(bib_path):
            bib_path.write_text(
                _generate_bibliography_stubs(all_claim_citation_keys),
                encoding="utf-8",
            )

    _pdf_ok, compiler = _compile_pdf_full(tex_path, output_dir)
    _cleanup_aux_files(output_dir)

    if compiler == "none":
        console.print(
            "[yellow]No LaTeX toolchain found. Generating DOCX instead...[/yellow]"
        )
        docx_path = _generate_docx(project, output_dir)
        console.print(
            f"[green]DOCX generated: {docx_path.relative_to(project_root)}[/green]"
        )
    else:
        if pdf_path.exists():
            console.print(f"[green]PDF generated: {rel_output}/paper.pdf[/green]")
            if not no_reveal:
                _reveal_output(pdf_path)
        else:
            console.print(
                f"[red]LaTeX compilation failed. Check {rel_output}/paper.log[/red]"
            )

    reports_dir = (
        output_dir.parent.parent / "reports"
        if output_dir.parent.name == "paper_generated"
        else output_dir / "reports"
    )
    reports_dir.mkdir(parents=True, exist_ok=True)

    # Run preflight & reports
    from paperforge.services.pdf_preflight import run_pdf_preflight
    from paperforge.services.reference_verifier import verify_references
    from paperforge.services.structural_integrity import check_structural_integrity
    from paperforge.services.template_fingerprint import verify_template_fingerprint

    fp_res = verify_template_fingerprint(latex, target, output_dir)
    struct_res = check_structural_integrity(
        project, reports_dir, mode=mode, tex_content=latex
    )
    ref_res = verify_references(project, reports_dir, online=False)

    pdf_preflight_passed = False
    overlap_passed = False
    artifact_passed = False

    if pdf_path.exists():
        pdf_res = run_pdf_preflight(pdf_path, reports_dir, mode=mode)
        pdf_preflight_passed = pdf_res.passed
        overlap_passed = not any(
            i.get("code") == "PDF_OBJECT_OVERLAP" and i.get("severity") == "ERROR"
            for i in pdf_res.issues
        )
        artifact_passed = not any(
            i.get("code") == "PDF_TEXT_ARTIFACT" and i.get("severity") == "ERROR"
            for i in pdf_res.issues
        )

    submission_ready = (
        _pdf_ok
        and pdf_preflight_passed
        and fp_res.passed
        and struct_res.passed
        and ref_res.passed
    )

    body_lines = [
        Text(
            f"LaTeX compilation:  {'PASSED ✓' if _pdf_ok else 'FAILED ✗'} ({compiler})"
        ),
        Text(f"PDF rendering:      {'PASSED ✓' if pdf_path.exists() else 'FAILED ✗'}"),
        Text(f"Visual overlap scan:{'PASSED ✓' if overlap_passed else 'FAILED ✗'}"),
        Text(f"Text artifact scan: {'PASSED ✓' if artifact_passed else 'FAILED ✗'}"),
        Text(f"Structural integrity:{'PASSED ✓' if struct_res.passed else 'FAILED ✗'}"),
        Text(f"Venue fingerprint:  {'PASSED ✓' if fp_res.passed else 'FAILED ✗'}"),
        Text(
            f"Submission readiness:{'PASSED ✓' if submission_ready else 'BLOCKED ✗'} ({mode} mode)"
        ),
        Text(""),
        Text(f"Output directory:   {rel_output}"),
        Text(f"Reports directory:  {reports_dir.relative_to(project_root)}"),
    ]

    unique_citations = {c for claim in project.claims for c in claim.citations}
    body = Group(
        *body_lines,
        Text(""),
        Text(f"Claims compiled:    {len(project.claims)}"),
        Text(f"Sections:           {len(project.config.sections)}"),
        Text(f"Citations:          {len(unique_citations)}"),
    )

    console.print(
        Panel(
            body,
            title="Build Summary",
            border_style="green" if submission_ready else "red",
        )
    )

    if venue_warnings:
        console.print()
        console.print(Text(f"VENUE ({plugin.display_name})", style="bold yellow"))
        for issue in venue_warnings:
            console.print(Text(f"  [{issue.code}] {issue.message}"))

    _generate_build_reports(project, issues + venue_issues, output_dir, mode)

    if submission_mode and not submission_ready and not force_anyway:
        console.print(
            "[red]Submission mode build blocked due to preflight/quality failures.[/red]"
        )
        sys.exit(1)
